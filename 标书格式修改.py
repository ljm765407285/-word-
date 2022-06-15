#!/usr/bin/env python
# _*_ coding: utf-8 _*_
#author: alber time:2022/6/10
import docx
from docx import Document
from docx.shared import Pt,RGBColor,Cm,Inches,Length # 字号，设置像素、缩进,颜色,宽度，厘米，英寸等
from docx.oxml.ns import qn # 中文字体
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_TAB_ALIGNMENT,WD_TAB_LEADER  #设置对象居中、对齐、制表符等
文件 = Document('d:/test/test2.docx')
j = int(len(文件.sections))
for i in range(j):
    节 = 文件.sections[i]
    节.page_width = Cm(21) #页面宽度
    节.page_height = Cm(29.7) #页面高度
    节.top_margin = Cm(2.54) #页边距上
    节.bottom_margin = Cm(2.54) #页边距下
    节.left_margin = Cm(1.91) #页边距左
    节.right_margin = Cm(1.91) #页边距右
    节.header_distance = Cm(1.5) #页眉距离1.5厘米
    节.footer_distance = Cm(1.5) #页脚距离1.5厘米
    页眉 = 节.header
    页眉段落 =页眉.paragraphs[0]
    #页眉段落.font.size = Pt(26)
    #页眉段落._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    页眉段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    页脚 = 节.footer
    页脚段落 = 页脚.paragraphs[0]
    #页脚段落.font.size = Pt(26)
    #页脚段落._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    页脚段落.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for 段落 in 文件.paragraphs:
    for 块 in 段落.runs:
        if 段落.style.name == 'Heading 1':
            块.font.size = Pt(22)
            块.font.name = 'Arial'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
            段落.paragraph_format.line_spacing = 1.5 #1.5倍行距
            段落.paragraph_format.space_before = Pt(0) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 2':
            块.font.size = Pt(16)
            块.font.name = 'Arial'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 3':
            块.font.size = Pt(15)
            块.font.name = 'Arial'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 3':
            块.font.size = Pt(14)
            块.font.name = 'Arial'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Heading 4':
            块.font.size = Pt(12)
            块.font.name = 'Arial'
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            块.font.color.rgb = RGBColor(0, 0, 0, )  # 颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0)  # 段前
            段落.paragraph_format.space_after = Pt(0)  # 段后
            段落.paragraph_format.left_indent = Inches(0)  # 正文前
            段落.paragraph_format.right_indent = Inches(0)  # 正文后
        elif 段落.style.name == 'Normal':
            块.font.size = Pt(10.5)
            块.font.name = 'Arial' # 设置英文字体
            块._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置中文字体
            块.font.color.rgb = RGBColor(0,0,0,) #颜色
            段落.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
            段落.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            段落.paragraph_format.space_before = Pt(0) #段前
            段落.paragraph_format.space_after = Pt(0) #段后
            段落.paragraph_format.left_indent = Inches(0) #正文前
            段落.paragraph_format.right_indent = Inches(0) #正文后
            段落.paragraph_format.first_line_indent = Cm(0.74) #首行缩进
文件.save('d:/test/test2.docx')